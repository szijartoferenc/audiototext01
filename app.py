import streamlit as st
import os
import torchaudio
import tempfile
from datetime import timedelta
from fpdf import FPDF
from docx import Document
from vosk import Model, KaldiRecognizer
import wave
import json
import re
import torch

st.set_page_config(page_title="🗣️ Hang → Szöveg konvertáló", layout="centered")

LANGUAGES = {
    "de": "Deutsch",
    "en": "English"
}

TEXTS = {
    "de": {
        "title": "🗣️ Audio → Text Konverter",
        "upload": "Nur WAV-Dateien hochladen",
        "processing": "⏳ Verarbeitung läuft...",
        "done": "✅ Verarbeitung abgeschlossen!",
        "recognized": "📝 Erkannter Text:",
        "export_docx": "📂 Als DOCX speichern",
        "export_pdf": "📄 Als PDF speichern",
        "export_srt": "🎬 Als SRT exportieren",
        "stats": "📊 Statistik",
        "words": "Wortanzahl",
        "duration": "Geschätzte Dauer",
        "model": "🧠 Modell auswählen",
        "theme": "🎨 Thema",
        "detected_lang": "🌍 Erkannte Sprache"
    },
    "en": {
        "title": "🗣️ Speech → Text Converter",
        "upload": "Upload WAV file only",
        "processing": "⏳ Processing...",
        "done": "✅ Done!",
        "recognized": "📝 Transcribed Text:",
        "export_docx": "📂 Save as DOCX",
        "export_pdf": "📄 Export PDF",
        "export_srt": "🎬 Export SRT",
        "stats": "📊 Statistics",
        "words": "Word Count",
        "duration": "Estimated Duration",
        "model": "🧠 Select Model",
        "theme": "🎨 Theme",
        "detected_lang": "🌍 Detected Language"
    }
}

lang = st.sidebar.selectbox("🌐 Nyelv / Language", list(LANGUAGES.keys()), format_func=lambda x: LANGUAGES[x])
TEXT = TEXTS[lang]

@st.cache_resource
def load_model(lang_code):
    model_paths = {
        "de": "models/vosk-model-small-de-0.15",
        "en": "models/vosk-model-small-en-us-0.15"
    }
    return Model(model_paths[lang_code])

model = load_model(lang)

def export_docx(text):
    doc = Document()
    doc.add_heading(TEXT["recognized"], 0)
    doc.add_paragraph(text)
    path = os.path.join(tempfile.gettempdir(), "output.docx")
    doc.save(path)
    return path

def export_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("helvetica", size=12)
    pdf.multi_cell(0, 10, text)
    path = os.path.join(tempfile.gettempdir(), "output.pdf")
    pdf.output(path)
    return path

def export_srt_from_words(word_segments, max_duration=10.0):
    def format_sentence(text):
        text = text.strip()
        if not text:
            return ""
        text = text[0].upper() + text[1:]
        if text[-1] not in ".!?":
            text += "."
        return text

    srt_output = ""
    sentence = ""
    start_time = None
    end_time = None
    i = 1

    for word_data in word_segments:
        word = word_data["word"]
        start = word_data["start"]
        end = word_data["end"]

        if start_time is None:
            start_time = start

        sentence += word + " "
        end_time = end

        if re.search(r"[.!?]$", word) or (end_time - start_time > max_duration):
            formatted = format_sentence(sentence)
            srt_output += f"{i}\n{str(timedelta(seconds=start_time))} --> {str(timedelta(seconds=end_time))}\n{formatted}\n\n"
            i += 1
            sentence = ""
            start_time = None
            end_time = None

    if sentence:
        formatted = format_sentence(sentence)
        srt_output += f"{i}\n{str(timedelta(seconds=start_time))} --> {str(timedelta(seconds=end_time))}\n{formatted}\n\n"

    path = os.path.join(tempfile.gettempdir(), "output.srt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(srt_output)
    return path

def load_and_resample_wav(wav_path, target_sample_rate=16000):
    waveform, sample_rate = torchaudio.load(wav_path)

    if waveform.ndimension() != 2 or waveform.shape[0] != 1:
        waveform = waveform.mean(dim=0, keepdim=True)

    if sample_rate != target_sample_rate:
        resampler = torchaudio.transforms.Resample(orig_freq=sample_rate, new_freq=target_sample_rate)
        waveform = resampler(waveform)
        sample_rate = target_sample_rate

    if waveform.dtype != torch.int16:
        waveform = (waveform * 32767).clamp(-32768, 32767).to(torch.int16)

    return waveform, sample_rate

st.title(TEXT["title"])
uploaded_file = st.file_uploader(TEXT["upload"], type=["wav"])

if uploaded_file is not None:
    if uploaded_file.size > 100 * 1024 * 1024:
        st.error("❌ A fájl túl nagy! Kérlek válassz egy kisebb WAV fájlt.")
    else:
        try:
            if uploaded_file.type != "audio/wav":
                st.error("❌ Csak WAV fájlokat fogadunk el!")
                st.stop()

            st.info(TEXT["processing"])

            with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
                tmp.write(uploaded_file.read())
                original_wav_path = tmp.name

            # Resample the audio
            waveform, sample_rate = load_and_resample_wav(original_wav_path)
            resampled_wav_path = os.path.join(tempfile.gettempdir(), "resampled.wav")
            torchaudio.save(resampled_wav_path, waveform, sample_rate)

            wf = wave.open(resampled_wav_path, "rb")

            if wf.getnchannels() != 1 or wf.getsampwidth() != 2:
                st.warning("⚠️ A fájl nem mono 16-bit WAV. Az eredmények pontatlanok lehetnek.")

            recognizer = KaldiRecognizer(model, sample_rate)
            recognizer.SetWords(True)

            result_text = ""
            segments = []

            while True:
                data = wf.readframes(4000)
                if len(data) == 0:
                    break
                if recognizer.AcceptWaveform(data):
                    res = json.loads(recognizer.Result())
                    result_text += res.get("text", "") + " "
                    if "result" in res:
                        segments.extend(res["result"])

            res = json.loads(recognizer.FinalResult())
            result_text += res.get("text", "")
            if "result" in res:
                segments.extend(res["result"])

            final_text = result_text.strip()

            st.success(TEXT["done"])
            st.subheader(TEXT["recognized"])
            if final_text:
                st.write(final_text)
            else:
                st.warning("⚠️ Nem sikerült értelmes szöveget felismerni.")

            st.subheader(TEXT["stats"])
            st.write(f"🌤️ {TEXT['words']}: {len(final_text.split())}")
            st.write(f"⏱️ {TEXT['duration']}: {timedelta(seconds=int(wf.getnframes() / sample_rate))}")

            col1, col2, col3 = st.columns(3)
            with col1:
                st.download_button(TEXT["export_docx"], open(export_docx(final_text), "rb"), file_name="output.docx")
            with col2:
                st.download_button(TEXT["export_pdf"], open(export_pdf(final_text), "rb"), file_name="output.pdf")
            with col3:
                st.download_button(TEXT["export_srt"], open(export_srt_from_words(segments), "rb"), file_name="output.srt")

            os.remove(original_wav_path)
            os.remove(resampled_wav_path)

        except Exception as e:
            st.error(f"❌ Hiba történt: {e}")

st.markdown("---")
st.caption("Powered by Vosk + Streamlit | Készítette Ferenc Szijarto")
